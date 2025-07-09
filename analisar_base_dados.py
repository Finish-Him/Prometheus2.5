"""
Script para analisar a estrutura da base de dados Oráculo 3.0
"""

import pandas as pd
import os
from docx import Document
import json

def analisar_excel(caminho_arquivo):
    """
    Analisa o arquivo Excel da base de dados Oráculo 3.0
    
    Args:
        caminho_arquivo: Caminho para o arquivo Excel
        
    Returns:
        Dicionário com informações sobre as planilhas
    """
    print(f"Analisando arquivo Excel: {caminho_arquivo}")
    
    # Carregar o arquivo Excel
    try:
        # Listar todas as planilhas sem carregar todo o conteúdo
        excel_file = pd.ExcelFile(caminho_arquivo)
        sheet_names = excel_file.sheet_names
        
        print(f"Planilhas encontradas: {len(sheet_names)}")
        for sheet in sheet_names:
            print(f"  - {sheet}")
        
        # Analisar cada planilha
        estrutura_planilhas = {}
        for sheet_name in sheet_names:
            print(f"\nAnalisando planilha: {sheet_name}")
            
            # Ler apenas as primeiras linhas para entender a estrutura
            df = pd.read_excel(caminho_arquivo, sheet_name=sheet_name, nrows=5)
            
            # Obter informações sobre as colunas
            colunas = list(df.columns)
            tipos_dados = {col: str(df[col].dtype) for col in colunas}
            
            # Adicionar ao dicionário de estrutura
            estrutura_planilhas[sheet_name] = {
                "colunas": colunas,
                "tipos_dados": tipos_dados,
                "num_linhas": len(pd.read_excel(caminho_arquivo, sheet_name=sheet_name, usecols=[0]))
            }
            
            print(f"  Número de colunas: {len(colunas)}")
            print(f"  Número de linhas: {estrutura_planilhas[sheet_name]['num_linhas']}")
            print(f"  Primeiras colunas: {colunas[:5] if len(colunas) > 5 else colunas}")
        
        return estrutura_planilhas
    
    except Exception as e:
        print(f"Erro ao analisar arquivo Excel: {str(e)}")
        return {}

def analisar_documentacao(caminho_arquivo):
    """
    Analisa o arquivo de documentação da base de dados
    
    Args:
        caminho_arquivo: Caminho para o arquivo de documentação
        
    Returns:
        Texto extraído do documento
    """
    print(f"Analisando documentação: {caminho_arquivo}")
    
    try:
        # Carregar o documento
        doc = Document(caminho_arquivo)
        
        # Extrair texto de parágrafos
        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        
        print(f"Número de parágrafos: {len(paragrafos)}")
        print("Primeiros parágrafos:")
        for i, p in enumerate(paragrafos[:5]):
            print(f"  {i+1}. {p[:100]}..." if len(p) > 100 else f"  {i+1}. {p}")
        
        # Extrair tabelas
        tabelas = []
        for i, table in enumerate(doc.tables):
            tabela_dados = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                tabela_dados.append(row_data)
            tabelas.append(tabela_dados)
        
        print(f"Número de tabelas: {len(tabelas)}")
        
        return {
            "paragrafos": paragrafos,
            "tabelas": tabelas
        }
    
    except Exception as e:
        print(f"Erro ao analisar documentação: {str(e)}")
        return {"paragrafos": [], "tabelas": []}

def analisar_liquipedia(caminho_arquivo):
    """
    Analisa o arquivo de análise da Liquipédia
    
    Args:
        caminho_arquivo: Caminho para o arquivo de análise
        
    Returns:
        Texto extraído do documento
    """
    print(f"Analisando documento Liquipédia: {caminho_arquivo}")
    
    try:
        # Carregar o documento
        doc = Document(caminho_arquivo)
        
        # Extrair texto de parágrafos
        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        
        print(f"Número de parágrafos: {len(paragrafos)}")
        print("Primeiros parágrafos:")
        for i, p in enumerate(paragrafos[:5]):
            print(f"  {i+1}. {p[:100]}..." if len(p) > 100 else f"  {i+1}. {p}")
        
        return {
            "paragrafos": paragrafos
        }
    
    except Exception as e:
        print(f"Erro ao analisar documento Liquipédia: {str(e)}")
        return {"paragrafos": []}

def salvar_resultados(resultados, caminho_saida):
    """
    Salva os resultados da análise em um arquivo JSON
    
    Args:
        resultados: Dicionário com resultados da análise
        caminho_saida: Caminho para o arquivo de saída
    """
    try:
        with open(caminho_saida, 'w', encoding='utf-8') as f:
            json.dump(resultados, f, ensure_ascii=False, indent=2)
        print(f"Resultados salvos em: {caminho_saida}")
    except Exception as e:
        print(f"Erro ao salvar resultados: {str(e)}")

def main():
    """Função principal"""
    diretorio_base = "/home/ubuntu/upload"
    
    # Caminhos para os arquivos
    caminho_excel = os.path.join(diretorio_base, "Database Oráculo 3.0.xlsx")
    caminho_doc = os.path.join(diretorio_base, "Documentação da Base de Dados Oráculo 3.docx")
    caminho_liquipedia = os.path.join(diretorio_base, "Análise detalhada da Liquipédia-Dota.docx")
    
    # Analisar os arquivos
    print("Iniciando análise da base de dados Oráculo 3.0...")
    
    # Analisar Excel
    estrutura_excel = analisar_excel(caminho_excel)
    
    # Analisar documentação
    conteudo_doc = analisar_documentacao(caminho_doc)
    
    # Analisar Liquipédia
    conteudo_liquipedia = analisar_liquipedia(caminho_liquipedia)
    
    # Combinar resultados
    resultados = {
        "estrutura_excel": estrutura_excel,
        "documentacao": conteudo_doc,
        "liquipedia": conteudo_liquipedia
    }
    
    # Salvar resultados
    salvar_resultados(resultados, "/home/ubuntu/analise_base_dados.json")
    
    print("Análise concluída!")

if __name__ == "__main__":
    main()
